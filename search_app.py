# ==============================================================================
# Universal Document Search Application
# ==============================================================================
#
# A desktop application with a GUI to perform semantic searches across
# local folders and Google Drive, ranking documents by relevance.
# It displays ranked results with matching sentences, line numbers, and page numbers.
# Now with full recursive search for Google Drive.
#

# --- 1. IMPORTS ---
import os
import io
import math
import re
import sys
import queue
from threading import Thread
import tkinter as tk
from tkinter import (
    Label, Button, Entry, OptionMenu, Listbox, MULTIPLE, StringVar, END,
    messagebox, scrolledtext, filedialog
)

# Third-party libraries
import numpy as np
import nltk
from nltk.corpus import stopwords
import PyPDF2
import openpyxl
import docx

# Google API libraries
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload

# --- 2. CONSTANTS AND INITIAL SETUP ---
try:
    nltk.download("stopwords", quiet=True)
    nltk.download("punkt", quiet=True)
    nltk.download("punkt_tab", quiet=True)
    STOP_WORDS = set(stopwords.words("english"))
except Exception as e:
    print(f"FATAL: Could not download NLTK data. Check internet connection. Error: {e}")
    sys.exit(1)

# --- 3. CORE LOGIC & FORMATTING FUNCTIONS ---

def preprocess_text(text):
    if not text: return []
    return [word.lower() for word in nltk.word_tokenize(text) if word.lower() not in STOP_WORDS and word.isalpha()]

def format_highlighted_line(line, query_terms):
    words = re.findall(r"(\w+)|(\W+)", line)
    highlighted = "".join([word.upper() if word and word.lower() in query_terms else (word or non_word) for word, non_word in words])
    return highlighted

def calculate_idf(processed_docs):
    all_tokens = set(word for doc in processed_docs for word in doc)
    vocab = sorted(list(all_tokens))
    num_docs = len(processed_docs)
    doc_freq = {term: sum(1 for doc in processed_docs if term in doc) for term in vocab}
    idf = {term: math.log(num_docs / (1 + doc_freq.get(term, 0))) for term in vocab}
    return vocab, idf

def create_tfidf_vector(text_tokens, vocab, idf):
    num_tokens = len(text_tokens)
    tf = {term: text_tokens.count(term) / num_tokens for term in vocab} if num_tokens > 0 else {}
    return np.array([tf.get(term, 0) * idf.get(term, 0) for term in vocab])

def cosine_similarity(v1, v2):
    dot = np.dot(v1, v2)
    norm = np.linalg.norm(v1) * np.linalg.norm(v2)
    return dot / norm if norm > 0 else 0.0

# --- 4. DATA LOADING FUNCTIONS ---

def load_documents_local(paths, update_queue):
    docs_content, doc_paths = [], []
    for path in paths:
        update_queue.put(f"--- Scanning Local Path: {path} ---")
        try:
            for root, _, files in os.walk(path):
                for file in files:
                    file_path, ext = os.path.join(root, file), file.split('.')[-1].lower()
                    update_queue.put(f"  -> Processing: {os.path.basename(file_path)}")
                    content = []
                    try:
                        if ext == 'pdf':
                            with open(file_path, 'rb') as f:
                                reader = PyPDF2.PdfReader(f)
                                for i, page in enumerate(reader.pages):
                                    if text := page.extract_text():
                                        for line in text.splitlines(True): content.append((i + 1, line))
                        elif ext in ['txt', 'docx', 'xlsx', 'xls']:
                            lines = []
                            if ext == 'txt':
                                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f: lines = f.readlines()
                            elif ext == 'docx':
                                doc = docx.Document(file_path)
                                lines = [p.text + '\\n' for p in doc.paragraphs if p.text]
                            elif ext in ['xlsx', 'xls']:
                                wb = openpyxl.load_workbook(file_path, data_only=True)
                                lines = [' '.join([str(c) for c in row if c])+'\\n' for s in wb.worksheets for row in s.iter_rows(values_only=True)]
                            for line in lines: content.append((None, line))
                        
                        if content: docs_content.append(content); doc_paths.append(file_path)
                    except Exception: continue
        except Exception as e: update_queue.put(f"ERROR: Could not access {path}. Reason: {e}")
    return docs_content, doc_paths

# **CORRECTED Google Drive loading function**
def load_documents_drive(service, folder_id, update_queue):
    docs_content, doc_paths = [], []

    def get_files_recursively(fid, current_path=""):
        """Helper to recursively get all files and their paths."""
        file_list = []
        page_token = None
        while True:
            response = service.files().list(q=f"'{fid}' in parents and trashed=false",
                                            fields="nextPageToken, files(id, name, mimeType)",
                                            pageToken=page_token).execute()
            for item in response.get('files', []):
                item_path = os.path.join(current_path, item['name'])
                if item['mimeType'] == 'application/vnd.google-apps.folder':
                    file_list.extend(get_files_recursively(item['id'], item_path))
                else:
                    # Store the file's ID and its full nested path
                    file_list.append({'id': item['id'], 'name': item_path, 'mimeType': item['mimeType']})
            page_token = response.get('nextPageToken')
            if not page_token: break
        return file_list

    update_queue.put("--- Scanning Google Drive ---")
    all_files = get_files_recursively(folder_id)

    for file in all_files:
        file_id, name, mime = file['id'], file['name'], file.get('mimeType', '')
        update_queue.put(f"  -> Processing: {name}")
        content = []
        try:
            # (The rest of the file processing logic is the same)
            if 'google-apps' in mime:
                export_map = {'application/vnd.google-apps.document': 'text/plain', 'application/vnd.google-apps.spreadsheet': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'}
                if mime not in export_map: continue
                request = service.files().export_media(fileId=file_id, mimeType=export_map[mime])
            else: request = service.files().get_media(fileId=file_id)
            
            fh = io.BytesIO(); downloader = MediaIoBaseDownload(fh, request); done = False
            while not done: _, done = downloader.next_chunk()
            fh.seek(0)
            
            ext = name.split('.')[-1].lower()
            if ext == 'pdf':
                reader = PyPDF2.PdfReader(fh)
                for i, page in enumerate(reader.pages):
                    if text := page.extract_text():
                        for line in text.splitlines(True): content.append((i + 1, line))
            else:
                lines = []
                if 'spreadsheet' in mime or ext in ['xlsx', 'xls']:
                    wb = openpyxl.load_workbook(fh, data_only=True)
                    lines = [' '.join([str(c) for c in row if c])+'\\n' for s in wb.worksheets for row in s.iter_rows(values_only=True)]
                elif 'document' in mime or ext == 'txt':
                    lines = fh.read().decode('utf-8', errors='ignore').splitlines(True)
                elif ext == 'docx':
                    doc = docx.Document(fh)
                    lines = [p.text + '\\n' for p in doc.paragraphs if p.text]
                for line in lines: content.append((None, line))
            
            if content: docs_content.append(content); doc_paths.append(name)
        except Exception: continue
    return docs_content, doc_paths

# --- 5. GUI APPLICATION CLASS ---

class SearchApp:
    # (This class remains identical to the previous final version)
    def __init__(self, master):
        self.master = master; master.title("Universal Document Search"); master.minsize(800, 650)
        self.master.grid_columnconfigure(1, weight=1)
        Label(master, text="1. Search Source:", font=("Helvetica", 10, "bold")).grid(row=0, column=0, sticky="w", padx=10, pady=5)
        self.source_var = StringVar(master, "Local Storage")
        OptionMenu(master, self.source_var, "Local Storage", "Google Drive", "Both", command=self.on_source_change).grid(row=0, column=1, sticky="w", padx=10)
        Label(master, text="Google Drive Folder ID:").grid(row=1, column=0, sticky="w", padx=10)
        self.gdrive_entry = Entry(master, width=60); self.gdrive_entry.grid(row=1, column=1, sticky="we", padx=10)
        Label(master, text="Local Folders:").grid(row=2, column=0, sticky="nw", padx=10, pady=(10, 0))
        self.local_listbox = Listbox(master, selectmode=MULTIPLE, height=5); self.local_listbox.grid(row=2, column=1, sticky="we", padx=10, pady=(10, 0))
        folder_buttons = tk.Frame(master); folder_buttons.grid(row=3, column=1, sticky="w", padx=10)
        self.add_folder_button = Button(folder_buttons, text="Add Folder...", command=self.add_folder); self.add_folder_button.pack(side="left", pady=3)
        self.clear_folders_button = Button(folder_buttons, text="Clear List", command=self.clear_folders); self.clear_folders_button.pack(side="left", padx=5, pady=3)
        Label(master, text="2. Search Query:", font=("Helvetica", 10, "bold")).grid(row=4, column=0, sticky="w", padx=10, pady=(10, 5))
        self.query_entry = Entry(master, width=60); self.query_entry.grid(row=4, column=1, sticky="we", padx=10)
        self.search_button = Button(master, text="Search", font=("Helvetica", 10, "bold"), command=self.start_search); self.search_button.grid(row=5, column=1, pady=10, padx=10, sticky="w")
        self.results_area = scrolledtext.ScrolledText(master, wrap=tk.WORD, state="disabled", font=("Courier New", 9)); self.results_area.grid(row=6, column=0, columnspan=2, padx=10, pady=10, sticky="nsew")
        self.master.grid_rowconfigure(6, weight=1); self.update_queue = queue.Queue(); self.on_source_change(); self.process_queue()

    def add_folder(self):
        folder = filedialog.askdirectory()
        if folder and folder not in self.local_listbox.get(0, END): self.local_listbox.insert(END, folder)
    def clear_folders(self): self.local_listbox.delete(0, END)
    def on_source_change(self, *args):
        source = self.source_var.get(); gdrive, local = "disabled", "disabled"
        if source in ["Google Drive", "Both"]: gdrive = "normal"
        if source in ["Local Storage", "Both"]: local = "normal"
        self.gdrive_entry.config(state=gdrive); self.local_listbox.config(state=local); self.add_folder_button.config(state=local); self.clear_folders_button.config(state=local)
        
    def process_queue(self):
        try:
            while True:
                msg = self.update_queue.get_nowait(); self.results_area.config(state="normal")
                if isinstance(msg, tuple) and msg[0] == "RESULT":
                    _, rank, score, path, match_tuples, query_terms = msg
                    self.results_area.insert(END, "─" * 80 + "\n")
                    self.results_area.insert(END, f"Rank {rank:<3} | Score: {score:.4f} | File: {path}\n\n", "bold")
                    if match_tuples:
                        for line_num, page_num, sent in match_tuples[:3]:
                            if not isinstance(sent, str): sent = str(sent)
                            highlighted = format_highlighted_line(sent, query_terms)
                            page_info = f"(Page: {page_num})" if page_num else ""
                            self.results_area.insert(END, f"Line {line_num:<4}: {highlighted} {page_info}\n")
                else: self.results_area.insert(END, msg + "\n")
                self.results_area.config(state="disabled"); self.results_area.see(END)
        except queue.Empty: pass
        self.master.after(100, self.process_queue)

    def start_search(self):
        self.results_area.config(state="normal"); self.results_area.delete('1.0', END); self.results_area.config(state="disabled")
        self.search_button.config(state="disabled"); Thread(target=self.perform_search, daemon=True).start()

    def perform_search(self):
        q = self.update_queue.put
        q("--- Scan & Ingestion Status ---")
        source, query = self.source_var.get(), self.query_entry.get().strip()
        if not query: messagebox.showerror("Input Error", "Query is required."); self.search_button.config(state="normal"); return
        docs_content, doc_paths = [], []
        if source in ["Google Drive", "Both"]:
            if not (gdrive_id := self.gdrive_entry.get().strip()): messagebox.showerror("Input Error", "Google Drive Folder ID required.")
            else:
                try: q("Authenticating Google Drive..."); service = self.authenticate_gdrive(); gdocs, gpaths = load_documents_drive(service, gdrive_id, self.update_queue); docs_content.extend(gdocs); doc_paths.extend([f"[GDrive] {p}" for p in gpaths])
                except Exception as e: q(f"Google Drive Error: {e}")
        if source in ["Local Storage", "Both"]:
            if not (local_paths := self.local_listbox.get(0, END)): messagebox.showerror("Input Error", "Please add a local folder.")
            else: ldocs, lfound = load_documents_local(local_paths, self.update_queue); docs_content.extend(ldocs); doc_paths.extend([f"[Local] {p}" for p in lfound])

        if not docs_content: q("\nNo documents found."); self.search_button.config(state="normal"); return
        
        q("\n--- Processing & Ranking ---")
        plain_docs_content = [[line_text for _, line_text in doc] for doc in docs_content]
        processed_query = preprocess_text(query)
        processed_docs = [preprocess_text(' '.join(doc)) for doc in plain_docs_content]
        vocab, idf = calculate_idf(processed_docs)
        query_vector = create_tfidf_vector(processed_query, vocab, idf)
        doc_vectors = [create_tfidf_vector(doc, vocab, idf) for doc in processed_docs]
        scores = [cosine_similarity(vec, query_vector) for vec in doc_vectors]

        results = []
        for i, score in enumerate(scores):
            if score > 0.001:
                match_tuples = []
                for line_idx, (page_num, line_text) in enumerate(docs_content[i]):
                    for sent in nltk.sent_tokenize(line_text):
                        if any(term in preprocess_text(sent) for term in processed_query):
                            match_tuples.append((line_idx + 1, page_num, sent.strip().replace('\n', ' ')))
                if match_tuples:
                    unique_matches = list({f"{ln},{pn},{s}": (ln, pn, s) for ln, pn, s in reversed(match_tuples)}.values())
                    unique_matches.reverse()
                    results.append({"path": doc_paths[i], "score": score, "matches": unique_matches})

        results.sort(key=lambda x: x["score"], reverse=True)
        
        q("\n--- Search Results ---")
        if not results: q("No relevant documents found.")
        else:
            for i, res in enumerate(results[:20]):
                q(("RESULT", i + 1, res['score'], res['path'], res['matches'], processed_query))
        self.search_button.config(state="normal")

    def authenticate_gdrive(self):
        SCOPES=['https://www.googleapis.com/auth/drive.readonly']; creds=None
        if os.path.exists('token.json'): creds=Credentials.from_authorized_user_file('token.json', SCOPES)
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token: creds.refresh(Request())
            else: flow=InstalledAppFlow.from_client_secrets_file('credentials.json',SCOPES); creds=flow.run_local_server(port=0)
            with open('token.json', 'w') as token: token.write(creds.to_json())
        return build('drive', 'v3', credentials=creds)

# --- 6. MAIN EXECUTION BLOCK ---
if __name__ == "__main__":
    root = tk.Tk()
    app = SearchApp(root)
    root.mainloop()

